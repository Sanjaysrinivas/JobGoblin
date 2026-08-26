"""Tests for plain-text extraction from PDF and DOCX (design.md §8)."""

import io

import pytest

from app.services.document_extractor import (
    SUPPORTED_CONTENT_TYPES,
    ExtractionError,
    UnsupportedDocumentError,
    extract_text,
)
from tests.sample_docs import (
    DOCX_CONTENT_TYPE,
    PDF_CONTENT_TYPE,
    make_docx_bytes,
    make_pdf_bytes,
)


def test_extract_pdf_text():
    data = make_pdf_bytes("Hello PDF\nSecond line")
    text = extract_text(data, PDF_CONTENT_TYPE)
    assert "Hello PDF" in text
    assert "Second line" in text


def test_extract_docx_text():
    data = make_docx_bytes("Hello DOCX\nAnother line")
    text = extract_text(data, DOCX_CONTENT_TYPE)
    assert "Hello DOCX" in text
    assert "Another line" in text


def test_extract_docx_includes_table_cells():
    from docx import Document

    document = Document()
    document.add_paragraph("Alice Example")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Skills"
    table.cell(0, 1).text = "Python Kubernetes SQL"
    buffer = io.BytesIO()
    document.save(buffer)

    text = extract_text(buffer.getvalue(), DOCX_CONTENT_TYPE)

    assert text.splitlines() == ["Alice Example", "Skills", "Python Kubernetes SQL"]


def test_unsupported_content_type_raises():
    with pytest.raises(UnsupportedDocumentError):
        extract_text(b"plain", "text/plain")


def test_legacy_doc_is_not_supported():
    # python-docx cannot read the legacy .doc binary format — we don't advertise it.
    assert "application/msword" not in SUPPORTED_CONTENT_TYPES
    with pytest.raises(UnsupportedDocumentError):
        extract_text(b"\xd0\xcf\x11\xe0", "application/msword")


def test_corrupt_pdf_raises_extraction_error():
    with pytest.raises(ExtractionError):
        extract_text(b"not a real pdf", PDF_CONTENT_TYPE)


def test_corrupt_docx_raises_extraction_error():
    with pytest.raises(ExtractionError):
        extract_text(b"not a real docx", DOCX_CONTENT_TYPE)
